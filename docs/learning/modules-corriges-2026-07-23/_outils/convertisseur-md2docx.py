#!/usr/bin/env python3
"""Convertit les modules corrigés (markdown) en .docx importables dans Google Docs.

Gère : titres, paragraphes, listes à puces et numérotées, tableaux, citations,
blocs de code, gras / italique / code inline, filets horizontaux.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def add_runs(paragraph, text):
    """Découpe le texte en fragments gras / italique / code et les ajoute."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def flush_table(doc, rows):
    """Écrit un tableau markdown accumulé."""
    rows = [r for r in rows if not re.fullmatch(r"\s*\|[\s:|-]+\|\s*", r)]
    if not rows:
        return
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    width = max(len(r) for r in cells)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        row += [""] * (width - len(row))
        cursor = table.add_row().cells
        for j, cell in enumerate(row):
            cursor[j].paragraphs[0].text = ""
            add_runs(cursor[j].paragraphs[0], cell)
            if i == 0:
                for run in cursor[j].paragraphs[0].runs:
                    run.bold = True


def convert(md_path: Path, docx_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").split("\n")
    table_buffer, code_buffer, in_code = [], [], False

    for line in lines:
        stripped = line.rstrip()

        # blocs de code
        if stripped.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                para.paragraph_format.left_indent = Pt(18)
                code_buffer, in_code = [], False
            else:
                in_code = True
            continue
        if in_code:
            code_buffer.append(stripped)
            continue

        # tableaux
        if stripped.startswith("|"):
            table_buffer.append(stripped)
            continue
        if table_buffer:
            flush_table(doc, table_buffer)
            table_buffer = []

        if not stripped:
            continue

        if re.fullmatch(r"-{3,}", stripped):
            doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # titres
        heading = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading:
            level = len(heading.group(1))
            para = doc.add_heading("", level=min(level, 4))
            add_runs(para, heading.group(2))
            continue

        # citations
        if stripped.startswith(">"):
            para = doc.add_paragraph(style="Intense Quote")
            add_runs(para, stripped.lstrip("> ").strip())
            continue

        # listes
        bullet = re.match(r"^\s*[-*]\s+(.*)", stripped)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            add_runs(para, bullet.group(1))
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.*)", stripped)
        if numbered:
            para = doc.add_paragraph(style="List Number")
            add_runs(para, numbered.group(1))
            continue

        add_runs(doc.add_paragraph(), stripped)

    if table_buffer:
        flush_table(doc, table_buffer)

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    for arg in sys.argv[1:]:
        src = Path(arg)
        out = src.with_suffix(".docx")
        convert(src, out)
        print(f"{src.name}  ->  {out.name}  ({out.stat().st_size // 1024} Ko)")
